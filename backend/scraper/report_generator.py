import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
import docx

# Register VML namespaces for horizontal vector shapes
nsmap['v'] = 'urn:schemas-microsoft-com:vml'
nsmap['o'] = 'urn:schemas-microsoft-com:office:office'

def add_horizontal_line(paragraph):
    """Adds a thin template-matching horizontal divider line to the paragraph using paragraph bottom border."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    else:
        bottom_existing = pBdr.find(qn('w:bottom'))
        if bottom_existing is not None:
            pBdr.remove(bottom_existing)
            
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')  # 1.5pt width
    bottom.set(qn('w:space'), '12')  # Spacing below the paragraph
    bottom.set(qn('w:color'), 'D3D3D3')  # Light grey
    pBdr.append(bottom)

def add_hyperlink(paragraph, url, text, color="1155cc", underline=True):
    """
    Inserts a clickable hyperlink run into a paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
        
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
    # Apply Calibri Font & Size 10pt
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rFonts)
    
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '20')  # 10pt
    rPr.append(sz)
    
    # Bold headline
    bold = OxmlElement('w:b')
    rPr.append(bold)
    
    new_run.append(rPr)
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def generate_docx_report(client_name: str, date_str: str, data: dict, output_path: str, template_path: str = None) -> str:
    """
    Generates a professionally formatted DOCX report for a client with outline headers & hyperlinked headlines.
    Matches the format of the provided reference Scapia Tracker.docx document.
    """
    # Ensure directory exists
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)

    if template_path and os.path.exists(template_path):
        doc = Document(template_path)
    else:
        doc = Document()
        
    # Set default margins if new document
    if not template_path:
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    # Determine default font name from template Normal style if possible
    default_font_name = 'Calibri'
    if template_path and os.path.exists(template_path):
        try:
            normal_font = doc.styles['Normal'].font.name
            if normal_font:
                default_font_name = normal_font
        except Exception:
            pass

    # 1. Add Logo at the top if it exists in the backend static folder
    logo_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "static", "logo.png")
    if os.path.exists(logo_path):
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo_p.paragraph_format.space_before = Pt(0)
        logo_p.paragraph_format.space_after = Pt(12)
        logo_run = logo_p.add_run()
        logo_run.add_picture(logo_path, width=docx.shared.Emu(1047750)) # Exact width from Scapia Tracker.docx (1.15 inches)

    # Format date to Option A: "17 June 2026"
    clean_date_str = date_str
    try:
        # Check if incoming format is "Month DD, YYYY" (e.g. "June 17, 2026")
        dt = datetime.strptime(date_str, "%B %d, %Y")
        clean_date_str = dt.strftime("%d %B %Y")
    except Exception:
        try:
            # Check if ISO date
            dt = datetime.fromisoformat(date_str)
            clean_date_str = dt.strftime("%d %B %Y")
        except Exception:
            pass

    # Document Header: Heading 1 style to automatically map to outline navigation tabs in Google Docs
    try:
        title_p = doc.add_paragraph(style='Heading 1')
    except (KeyError, ValueError):
        # Fallback if the custom template lacks a 'Heading 1' style definition
        title_p = doc.add_paragraph()
    title_run = title_p.add_run(clean_date_str)
    title_run.font.name = default_font_name
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    if not template_path:
        title_run.font.color.rgb = RGBColor(0, 0, 0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(12)

    # Write each section
    for section_name, articles in data.items():
        # Section Heading styled manually (not Heading style) so it doesn't map to outline nested under the date
        section_p = doc.add_paragraph()
        section_run = section_p.add_run(section_name)
        section_run.font.name = default_font_name
        section_run.font.size = Pt(10)
        section_run.font.bold = True
        if not template_path:
            section_run.font.color.rgb = RGBColor(0, 0, 0)
        section_p.paragraph_format.space_before = Pt(12)
        section_p.paragraph_format.space_after = Pt(12)

        if not articles:
            no_art_p = doc.add_paragraph()
            no_art_run = no_art_p.add_run("No relevant articles found for this section.")
            no_art_run.font.name = default_font_name
            no_art_run.font.size = Pt(10)
            no_art_run.font.italic = True
            if not template_path:
                no_art_run.font.color.rgb = RGBColor(0, 0, 0)
            no_art_p.paragraph_format.space_before = Pt(12)
            no_art_p.paragraph_format.space_after = Pt(12)
            continue

        # Group articles by category (A, B, or C)
        from scraper.search_utils import match_publication_category
        grouped_articles = {"A": [], "B": [], "C": []}
        for article in articles:
            cat = article.get("publication_category")
            if not cat:
                cat = match_publication_category(article.get("agency"), article.get("url") or article.get("resolved_url"))
            if cat not in ["A", "B", "C"]:
                cat = "C"
            grouped_articles[cat].append(article)

        # Write grouped categories
        for cat_name in ["A", "B", "C"]:
            cat_list = grouped_articles[cat_name]
            if not cat_list:
                continue

            # Sub-heading for Category
            cat_p = doc.add_paragraph()
            cat_run = cat_p.add_run(f"Category {cat_name} Publications")
            cat_run.font.name = default_font_name
            cat_run.font.size = Pt(10)
            cat_run.font.bold = True
            cat_run.font.italic = True
            cat_run.font.color.rgb = RGBColor(120, 120, 120)
            cat_p.paragraph_format.space_before = Pt(12)
            cat_p.paragraph_format.space_after = Pt(12)

            for i, article in enumerate(cat_list):
                title = article.get("title", "No Title")
                agency = article.get("agency") or "Unknown Publication"
                summary = article.get("summary") or "No summary available."
                url = article.get("url") or article.get("resolved_url") or "#"
                
                # Single paragraph for the entire article block (Calibri 10pt)
                art_p = doc.add_paragraph()
                art_p.paragraph_format.space_before = Pt(12)
                art_p.paragraph_format.space_after = Pt(12)
                art_p.paragraph_format.line_spacing = 1.15
                
                # 1. Headline (Hyperlink)
                try:
                    add_hyperlink(art_p, url, title, color="1155cc", underline=True)
                except Exception:
                    # Fallback to plain text if hyperlink XML injection fails
                    fallback_run = art_p.add_run(title)
                    fallback_run.font.name = default_font_name
                    fallback_run.font.size = Pt(10)
                    fallback_run.font.bold = True
                    fallback_run.font.color.rgb = RGBColor(17, 85, 204)
                
                # Paywall tag (if applicable)
                if article.get("is_paywalled"):
                    pw_run = art_p.add_run("  🔒 Paywalled")
                    pw_run.font.name = default_font_name
                    pw_run.font.size = Pt(9)
                    pw_run.font.bold = True
                    pw_run.font.color.rgb = RGBColor(204, 102, 0)
                
                # Publication name on same line (Bold, black, separated by " - ")
                pub_run = art_p.add_run(f" - {agency}\n\n")
                pub_run.font.name = default_font_name
                pub_run.font.size = Pt(10)
                pub_run.font.bold = True
                if not template_path:
                    pub_run.font.color.rgb = RGBColor(0, 0, 0)
                
                # 2. Summary Content (Second line onwards)
                sum_run = art_p.add_run(summary)
                sum_run.font.name = default_font_name
                sum_run.font.size = Pt(10)
                if not template_path:
                    sum_run.font.color.rgb = RGBColor(0, 0, 0)
                
                # Thin divider line between articles within the category
                if i < len(cat_list) - 1:
                    add_horizontal_line(art_p)
                
    # Save the document
    doc.save(output_path)
    return output_path


def generate_organized_docx_report(client_name: str, report_type: str, date_str: str, grouped_data: dict, output_path: str) -> str:
    """
    Generates a professionally formatted DOCX report grouped by Master Heading and Sub Heading.
    """
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(f"NEXUS NEWS BRIEFING: {report_type.upper()}")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(74, 134, 232)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add Subtitle / Date
    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run(f"Date: {date_str} | Generated for {client_name}")
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Divider line
    divider = doc.add_paragraph()
    add_horizontal_line(divider)

    if not grouped_data:
        empty_p = doc.add_paragraph()
        empty_run = empty_p.add_run("No relevant articles found for this briefing.")
        empty_run.font.name = 'Calibri'
        empty_run.font.size = Pt(11)
        doc.save(output_path)
        return output_path

    # Iterate over Master Headings
    for master, subs in grouped_data.items():
        # Add Master Heading (H1)
        m_p = doc.add_paragraph()
        m_run = m_p.add_run(master)
        m_run.font.name = 'Calibri'
        m_run.font.size = Pt(14)
        m_run.font.bold = True
        m_run.font.color.rgb = RGBColor(74, 134, 232)
        m_p.paragraph_format.space_before = Pt(18)
        m_p.paragraph_format.space_after = Pt(6)
        m_p.paragraph_format.keep_with_next = True

        for sub, articles in subs.items():
            if not articles:
                continue
            # Add Sub Heading (H2)
            s_p = doc.add_paragraph()
            s_run = s_p.add_run(sub)
            s_run.font.name = 'Calibri'
            s_run.font.size = Pt(12)
            s_run.font.bold = True
            s_run.font.color.rgb = RGBColor(102, 102, 102)
            s_p.paragraph_format.space_before = Pt(12)
            s_p.paragraph_format.space_after = Pt(4)
            s_p.paragraph_format.keep_with_next = True

            for i, art in enumerate(articles):
                # Add Article Headline as clickable link
                art_p = doc.add_paragraph()
                art_p.paragraph_format.space_before = Pt(6)
                art_p.paragraph_format.space_after = Pt(2)
                
                title_text = art.get("title", "Untitled Article")
                url = art.get("url", "#")
                
                try:
                    add_hyperlink(art_p, url, title_text, color="1155cc", underline=True)
                except Exception:
                    fallback_run = art_p.add_run(title_text)
                    fallback_run.font.name = 'Calibri'
                    fallback_run.font.size = Pt(10)
                    fallback_run.font.bold = True
                    fallback_run.font.color.rgb = RGBColor(17, 85, 204)
                
                # Add Publication, Author (if available), and Date
                meta_p = doc.add_paragraph()
                meta_p.paragraph_format.space_before = Pt(0)
                meta_p.paragraph_format.space_after = Pt(8)
                
                pub_name = art.get("agency") or "Unknown Publication"
                pub_author = art.get("author") or art.get("journalist")
                if not pub_author or str(pub_author).strip().upper() in ("N/A", "NONE", "NULL", ""):
                    pub_author = f"syndicated by the {pub_name}"
                
                pub_date = art.get("published_at")
                
                date_text = ""
                if pub_date:
                    if isinstance(pub_date, str):
                        date_text = pub_date[:10]
                    else:
                        date_text = pub_date.strftime("%Y-%m-%d")
                
                conf_score = art.get("confidence_score", 0)
                meta_text = f"Publication: {pub_name} | Author: {str(pub_author).strip()} | Date: {date_text} | Relevance Confidence: {conf_score}/10"
                
                meta_run = meta_p.add_run(meta_text)
                meta_run.font.name = 'Calibri'
                meta_run.font.size = Pt(9.5)
                meta_run.font.italic = True
                meta_run.font.color.rgb = RGBColor(128, 128, 128)
                
                # Thin divider line between articles
                if i < len(articles) - 1:
                    add_horizontal_line(meta_p)
                
    doc.save(output_path)
    return output_path


def generate_excel_report(client_name: str, report_type: str, date_str: str, grouped_data: dict, output_path: str) -> str:
    """
    Generates a clean, section-based Excel report (.xlsx) for a client run.
    Structures headings/sub-headings as full-width divider rows rather than repeating columns.
    Columns: Link, Title, Author, Publication, Date, Relevance Score.
    """
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Briefing Report"

    # Style definitions (Steel Blue & Ice Blue theme)
    master_fill = PatternFill(start_color="365F91", end_color="365F91", fill_type="solid")
    master_font = Font(name="Calibri", size=12, bold=True, color="FFFFFF")
    
    sub_fill = PatternFill(start_color="DCE6F1", end_color="DCE6F1", fill_type="solid")
    sub_font = Font(name="Calibri", size=11, bold=True, color="1E3A5F")
    
    table_header_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
    table_header_font = Font(name="Calibri", size=10, bold=True, color="000000")
    
    data_font = Font(name="Calibri", size=10)
    link_font = Font(name="Calibri", size=10, color="0000FF", underline="single")
    title_font = Font(name="Calibri", size=14, bold=True, color="365F91")
    subtitle_font = Font(name="Calibri", size=11, italic=True)

    thin_border = Border(
        left=Side(style='thin', color='D3D3D3'),
        right=Side(style='thin', color='D3D3D3'),
        top=Side(style='thin', color='D3D3D3'),
        bottom=Side(style='thin', color='D3D3D3')
    )

    # 1. Title Block
    ws.cell(row=1, column=1, value=f"NEXUS NEWS BRIEFING: {report_type.upper()}").font = title_font
    ws.cell(row=2, column=1, value=f"Date: {date_str} | Generated for {client_name}").font = subtitle_font
    ws.row_dimensions[1].height = 25
    ws.row_dimensions[2].height = 20

    current_row = 4

    headers = [
        "Link of the article",
        "Title of the article",
        "Author of the article",
        "Name of publication",
        "Time of publishing",
        "Relevance score",
        "Summary of the article"
    ]

    # 2. Iterate through sections
    for master, subs in grouped_data.items():
        # Check if there are any articles in this master section
        has_articles_in_master = any(len(articles) > 0 for articles in subs.values())
        if not has_articles_in_master:
            continue

        # Add Master Heading Row (Merged A to G)
        ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=7)
        c_master = ws.cell(row=current_row, column=1, value=master)
        c_master.font = master_font
        c_master.fill = master_fill
        c_master.alignment = Alignment(vertical="center", indent=1)
        ws.row_dimensions[current_row].height = 26
        current_row += 1

        for sub, articles in subs.items():
            if not articles:
                continue

            # Add Sub Heading Row (Merged A to G)
            ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=7)
            c_sub = ws.cell(row=current_row, column=1, value=sub)
            c_sub.font = sub_font
            c_sub.fill = sub_fill
            c_sub.alignment = Alignment(vertical="center", indent=2)
            ws.row_dimensions[current_row].height = 22
            current_row += 1

            # Add Table Headers for the nested table
            for col_idx, h in enumerate(headers, start=1):
                cell = ws.cell(row=current_row, column=col_idx, value=h)
                cell.fill = table_header_fill
                cell.font = table_header_font
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                cell.border = thin_border
            ws.row_dimensions[current_row].height = 20
            current_row += 1

            # Add articles
            for art in articles:
                title = art.get("title", "Untitled Article")
                url = art.get("url") or art.get("resolved_url") or "#"
                author = art.get("author") or "N/A"
                pub_name = art.get("agency") or "Unknown Publication"
                pub_date = art.get("published_at")
                
                date_text = "N/A"
                if pub_date:
                    if isinstance(pub_date, str):
                        date_text = pub_date[:10]
                    else:
                        date_text = pub_date.strftime("%Y-%m-%d")
                        
                conf_score = art.get("confidence_score", 0)
                score_text = f"{conf_score}/10"

                # Write values
                c_link = ws.cell(row=current_row, column=1, value="Link")
                c_link.hyperlink = url
                c_link.font = link_font
                c_link.alignment = Alignment(horizontal="center")
                
                c_title = ws.cell(row=current_row, column=2, value=title)
                
                # Author fallback check
                author_val = str(author).strip()
                if not author_val or author_val.upper() in ("N/A", "NONE", "NULL", ""):
                    author_val = f"syndicated by the {pub_name}"
                c_author = ws.cell(row=current_row, column=3, value=author_val)
                
                c_pub = ws.cell(row=current_row, column=4, value=pub_name)
                c_date = ws.cell(row=current_row, column=5, value=date_text)
                
                c_score = ws.cell(row=current_row, column=6, value=score_text)
                c_score.alignment = Alignment(horizontal="center")
                
                summary_text = art.get("summary") or art.get("_summary") or art.get("full_body") or ""
                c_summary = ws.cell(row=current_row, column=7, value=summary_text)

                # Borders and alignment
                for col_idx in range(1, 8):
                    c = ws.cell(row=current_row, column=col_idx)
                    c.border = thin_border
                    if col_idx == 1:
                        pass
                    elif col_idx == 6:
                        c.font = data_font
                    elif col_idx == 7:
                        c.font = data_font
                        c.alignment = Alignment(vertical="center", wrap_text=True)
                    else:
                        c.font = data_font
                        c.alignment = Alignment(vertical="center", wrap_text=False)

                ws.row_dimensions[current_row].height = 20
                current_row += 1

            # Insert an empty row for visual breathing space after table
            current_row += 1

    # 3. Set specific widths for readability
    ws.column_dimensions["A"].width = 15  # Link
    ws.column_dimensions["B"].width = 50  # Title
    ws.column_dimensions["C"].width = 25  # Author
    ws.column_dimensions["D"].width = 25  # Publication Name
    ws.column_dimensions["E"].width = 15  # Date
    ws.column_dimensions["F"].width = 15  # Relevance Score
    ws.column_dimensions["G"].width = 60  # Summary of the article

    wb.save(output_path)
    return output_path


def generate_mailer_docx_report(client_name: str, report_type: str, date_str: str, exec_summary: str, takeaways: str, grouped_data: dict, output_path: str) -> str:
    """
    Generates a professionally formatted DOCX report representing the entire briefing mailer:
    Title -> Executive Summary -> Strategic Takeaways -> Grouped Articles.
    """
    import re
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(f"NEXUS NEWS BRIEFING: {report_type.upper()}")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(74, 134, 232)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle / Date
    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run(f"Date: {date_str} | Generated for {client_name}")
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Divider line
    divider = doc.add_paragraph()
    add_horizontal_line(divider)

    # 1. Add Executive Summary
    if exec_summary:
        es_h = doc.add_paragraph()
        es_hrun = es_h.add_run("EXECUTIVE SUMMARY")
        es_hrun.font.name = 'Calibri'
        es_hrun.font.size = Pt(14)
        es_hrun.font.bold = True
        es_hrun.font.color.rgb = RGBColor(74, 134, 232)
        es_h.paragraph_format.space_before = Pt(14)
        es_h.paragraph_format.space_after = Pt(6)

        # Parse and write lines
        for line in exec_summary.split("\n"):
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            if line.isupper() and len(line) < 30:
                # Label!
                run = p.add_run(line)
                run.font.name = 'Calibri'
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(66, 133, 244) # Blue accent
            else:
                run = p.add_run(line)
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)

        # Divider
        divider2 = doc.add_paragraph()
        add_horizontal_line(divider2)

    # 2. Add Strategic Takeaways
    if takeaways:
        st_h = doc.add_paragraph()
        st_hrun = st_h.add_run("STRATEGIC TAKEAWAYS")
        st_hrun.font.name = 'Calibri'
        st_hrun.font.size = Pt(14)
        st_hrun.font.bold = True
        st_hrun.font.color.rgb = RGBColor(74, 134, 232)
        st_h.paragraph_format.space_before = Pt(14)
        st_h.paragraph_format.space_after = Pt(6)

        for line in takeaways.split("\n"):
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line)
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)

        # Divider
        divider3 = doc.add_paragraph()
        add_horizontal_line(divider3)

    # 3. Add Grouped Categories & Articles
    if grouped_data:
        for master, subs in grouped_data.items():
            # Add Master Heading (H1)
            m_p = doc.add_paragraph()
            m_run = m_p.add_run(master)
            m_run.font.name = 'Calibri'
            m_run.font.size = Pt(14)
            m_run.font.bold = True
            m_run.font.color.rgb = RGBColor(74, 134, 232)
            m_p.paragraph_format.space_before = Pt(18)
            m_p.paragraph_format.space_after = Pt(6)
            m_p.paragraph_format.keep_with_next = True

            for sub, articles in subs.items():
                if not articles:
                    continue
                # Add Sub Heading (H2)
                s_p = doc.add_paragraph()
                s_run = s_p.add_run(sub)
                s_run.font.name = 'Calibri'
                s_run.font.size = Pt(12)
                s_run.font.bold = True
                s_run.font.color.rgb = RGBColor(102, 102, 102)
                s_p.paragraph_format.space_before = Pt(12)
                s_p.paragraph_format.space_after = Pt(4)
                s_p.paragraph_format.keep_with_next = True

                for art in articles:
                    # Add Headline link
                    art_p = doc.add_paragraph()
                    art_p.paragraph_format.space_before = Pt(6)
                    art_p.paragraph_format.space_after = Pt(2)
                    title_text = art.get("title", "Untitled Article")
                    url = art.get("url", "#")
                    try:
                        add_hyperlink(art_p, url, title_text, color="1155cc", underline=True)
                    except Exception:
                        fallback_run = art_p.add_run(title_text)
                        fallback_run.font.name = 'Calibri'
                        fallback_run.font.size = Pt(10)
                        fallback_run.font.bold = True
                        fallback_run.font.color.rgb = RGBColor(17, 85, 204)

                    # Add Publication & Journalist Byline
                    meta_p = doc.add_paragraph()
                    meta_p.paragraph_format.space_before = Pt(0)
                    meta_p.paragraph_format.space_after = Pt(6)
                    
                    pub_name = art.get("agency") or "Unknown Publication"
                    pub_author = art.get("author") or art.get("journalist")
                    if pub_author and str(pub_author).strip().upper() not in ("N/A", "NONE", "NULL", ""):
                        byline_str = f"{pub_name} | {pub_author}"
                    else:
                        byline_str = f"syndicated by the {pub_name}"
                        
                    meta_run = meta_p.add_run(byline_str)
                    meta_run.font.name = 'Calibri'
                    meta_run.font.size = Pt(9.5)
                    meta_run.font.color.rgb = RGBColor(128, 128, 128)

                    # Add Summary
                    sum_p = doc.add_paragraph()
                    sum_p.paragraph_format.space_before = Pt(0)
                    sum_p.paragraph_format.space_after = Pt(10)
                    sum_text = art.get("summary") or art.get("full_body") or ""
                    sum_run = sum_p.add_run(sum_text)
                    sum_run.font.name = 'Calibri'
                    sum_run.font.size = Pt(10)
                    sum_run.font.color.rgb = RGBColor(60, 60, 60)

    doc.save(output_path)
    return output_path

