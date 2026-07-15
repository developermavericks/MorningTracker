import os
import pytest
from docx import Document
from docx.shared import Pt, Inches
from docx.dml.color import RGBColor

# Import functions to test
from utils.google_docs import merge_docx_files
from scraper.report_generator import generate_docx_report

def test_case_insensitive_extension():
    # Test helper mimicking backend/routers/clients.py check
    valid_names = ["template.docx", "TEMPLATE.DOCX", "my_theme.Docx"]
    invalid_names = ["template.doc", "template.txt", "template.docx.zip"]
    
    for name in valid_names:
        assert name.lower().endswith(".docx")
        
    for name in invalid_names:
        assert not name.lower().endswith(".docx")

def test_relative_path_resolution():
    # Test path resolution logic
    db_stored_paths = [
        "client_1_template.docx",
        "e:\\MAVERICKS\\zMorning_Tracker_Synced_Git\\backend\\templates\\client_1_template.docx",
        "/app/backend/templates/client_1_template.docx"
    ]
    
    backend_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    templates_dir = os.path.join(backend_dir, "templates")
    
    for path in db_stored_paths:
        filename = os.path.basename(path)
        resolved = os.path.join(templates_dir, filename)
        assert resolved.endswith(os.path.join("templates", "client_1_template.docx"))

def test_style_inheritance_from_template(tmp_path):
    # 1. Create a dummy template docx with custom styles
    template_file = os.path.join(tmp_path, "test_template.docx")
    doc = Document()
    normal_style = doc.styles['Normal']
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(12)
    doc.save(template_file)
    
    # 2. Generate report using this template
    report_file = os.path.join(tmp_path, "output_report.docx")
    data = {"Section 1": [{"title": "Test Title", "agency": "Test Agency", "summary": "Test Summary", "url": "http://test.com"}]}
    
    generate_docx_report(
        client_name="Test Client",
        date_str="2026-07-02",
        data=data,
        output_path=report_file,
        template_path=template_file
    )
    
    # 3. Read report and check font inherited from template
    report_doc = Document(report_file)
    # Check that added run uses dynamic font
    found_arial = False
    for p in report_doc.paragraphs:
        for r in p.runs:
            if r.font.name == "Arial":
                found_arial = True
                
    assert found_arial, "Run did not inherit font 'Arial' from template"

def test_duplicate_paragraphs_merging(tmp_path):
    # 1. Create template with 2 template paragraphs (e.g. Header + Logo placeholder)
    template_file = os.path.join(tmp_path, "template.docx")
    temp_doc = Document()
    temp_doc.add_paragraph("Template Header")
    temp_doc.add_paragraph("Template Logo Placeholder")
    temp_doc.save(template_file)
    
    # 2. Create daily briefing 1 using generate_docx_report (which clears template text)
    daily_file = os.path.join(tmp_path, "daily_briefing.docx")
    daily_data = {"Section A": [{"title": "Today's News Article 1", "url": "http://test.com", "summary": "Summary 1"}]}
    generate_docx_report(
        client_name="Test Client",
        date_str="2026-07-15",
        data=daily_data,
        output_path=daily_file,
        template_path=template_file
    )
    
    # 3. Create existing monthly file using generate_docx_report
    monthly_file = os.path.join(tmp_path, "monthly_archive.docx")
    monthly_data = {"Section A": [{"title": "Yesterday's News Article 2", "url": "http://test.com", "summary": "Summary 2"}]}
    generate_docx_report(
        client_name="Test Client",
        date_str="2026-07-14",
        data=monthly_data,
        output_path=monthly_file,
        template_path=template_file
    )
    
    # 4. Merge them using merge_docx_files, passing the template_path
    combined_file = os.path.join(tmp_path, "combined.docx")
    merge_docx_files(daily_file, monthly_file, combined_file, template_path=template_file)
    
    # 5. Read combined and check paragraph structure
    combined_doc = Document(combined_file)
    paragraphs = [p.text for p in combined_doc.paragraphs if p.text.strip()]
    
    # Check that "Template Header" and "Template Logo Placeholder" are stripped, but articles are present
    assert not any("Template Header" in p for p in paragraphs)
    assert not any("Template Logo Placeholder" in p for p in paragraphs)
    assert any("Today's News Article 1" in p for p in paragraphs)
    assert any("Yesterday's News Article 2" in p for p in paragraphs)
