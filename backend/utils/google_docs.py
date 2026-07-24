import os
import logging
from typing import Any, Optional
from datetime import datetime, date
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
import docx
from docx import Document
from docx.text.run import Run
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap

# Register VML namespaces for horizontal vector shapes
nsmap['v'] = 'urn:schemas-microsoft-com:vml'
nsmap['o'] = 'urn:schemas-microsoft-com:office:office'

logger = logging.getLogger(__name__)

SCOPES = [
    'https://www.googleapis.com/auth/drive',
    'https://www.googleapis.com/auth/documents'
]

def add_horizontal_line(paragraph):
    """Adds a thin template-matching horizontal divider line to the paragraph using paragraph bottom border."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    else:
        bottom_existing = pBdr.find(qn('w:bottom'))
        if bottom_existing is not None:
            pBdr.remove(bottom_existing)
            
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')  # 1.5pt width
    bottom.set(qn('w:space'), '12')  # Spacing below the paragraph
    bottom.set(qn('w:color'), 'D3D3D3')  # Light grey
    pBdr.append(bottom)

def add_hyperlink(paragraph, url, text, color="1155cc", underline=True):
    """Adds a hyperlink with custom style (Calibri, blue, bold, underline) to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
        
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rFonts)
    
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '20')  # 10pt
    rPr.append(sz)
    
    bold = OxmlElement('w:b')
    rPr.append(bold)
    
    new_run.append(rPr)
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def merge_docx_files(new_docx_path: str, existing_docx_path: str, output_path: str, template_path: str = None):
    """
    Combines two DOCX files.
    The new daily briefing is prepended at the top, followed by a divider, 
    and then the previous days' content from the existing monthly file.
    Uses new_docx_path as the base to preserve images (logo) and styles perfectly,
    and triages standard runs and hyperlink runs from existing_docx_path.
    """
    # Load new daily briefing as the base
    combined = Document(new_docx_path)
    # Load existing monthly doc
    existing_doc = Document(existing_docx_path)
    
    # Check template paragraph count to avoid duplicate template headers/footers/logos
    N_temp = 0
    if template_path and os.path.exists(template_path):
        try:
            temp_doc = Document(template_path)
            # Safely clear body elements except sectPr to get the correct cleared paragraph count (which should be 0)
            body = temp_doc.element.body
            for child in list(body):
                if child.tag.endswith('sectPr'):
                    continue
                body.remove(child)
            N_temp = len(temp_doc.paragraphs)
        except Exception as e:
            logger.error(f"Failed to read template paragraphs: {e}")

    # 1. Add page break / separator at the end of combined
    sep_p = combined.add_paragraph()
    sep_run = sep_p.add_run("\n" + "="*40 + "\n")
    sep_run.font.color.rgb = RGBColor(180, 180, 180)
    sep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 2. Copy paragraphs from existing_doc to the end of combined
    paragraphs_to_copy = existing_doc.paragraphs[N_temp:]
    for paragraph in paragraphs_to_copy:
        new_p = combined.add_paragraph()
        if paragraph.style:
            try:
                new_p.style = paragraph.style
            except Exception:
                pass
        new_p.alignment = paragraph.alignment
        new_p.paragraph_format.space_before = paragraph.paragraph_format.space_before
        new_p.paragraph_format.space_after = paragraph.paragraph_format.space_after
        new_p.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing
        
        # Iterate over child elements of paragraph XML block to copy standard runs and hyperlinks
        for child in paragraph._p:
            tag_name = child.tag
            if tag_name.endswith('hyperlink'):
                # Hyperlink node!
                rId = child.get(qn('r:id'))
                if rId and rId in paragraph.part.rels:
                    url = paragraph.part.rels[rId].target_ref
                    # Extract text inside hyperlink
                    link_text = ""
                    for r_child in child:
                        if r_child.tag.endswith('r'):
                            r_obj = Run(r_child, paragraph)
                            link_text += r_obj.text
                    
                    # Recreate hyperlink in target
                    add_hyperlink(new_p, url, link_text, color="1155cc", underline=True)
            elif tag_name.endswith('r'):
                # Check if run contains drawing or pict elements (e.g. divider lines)
                has_pict = child.find(qn('w:pict')) is not None or child.find(qn('w:drawing')) is not None
                if has_pict:
                    import copy
                    copied_r = copy.deepcopy(child)
                    new_p._p.append(copied_r)
                else:
                    run = Run(child, paragraph)
                    if run.text:
                        new_run = new_p.add_run(run.text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        new_run.font.name = run.font.name
                        new_run.font.size = run.font.size
                        if run.font.color and run.font.color.rgb:
                            new_run.font.color.rgb = run.font.color.rgb
                        
        if paragraph._p.pPr is not None and paragraph._p.pPr.find(qn('w:pBdr')) is not None:
            add_horizontal_line(new_p)
            
    combined.save(output_path)

def get_drive_service():
    """Initializes the Google Drive API service using environment variable or local credentials JSON."""
    import json
    
    # 1. Try to load credentials from environment variable first (best for Railway/production)
    creds_json = os.environ.get("GOOGLE_CREDENTIALS_JSON")
    if creds_json:
        try:
            creds_json = creds_json.strip("'\"")
            creds_data = json.loads(creds_json)
            if "private_key" in creds_data:
                creds_data["private_key"] = creds_data["private_key"].replace("\\n", "\n")
            creds = service_account.Credentials.from_service_account_info(creds_data, scopes=SCOPES)
            return build('drive', 'v3', credentials=creds, cache_discovery=False)
        except Exception as e:
            logger.error(f"Failed to load Google credentials from environment variable: {e}")
            
    # 2. Fallback to local file (development)
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    creds_path = os.path.join(base_dir, "google_credentials.json")
    
    if not os.path.exists(creds_path):
        logger.warning(f"Google credentials file not found at {creds_path} and no environment variable set. Skipping Google Docs integration.")
        return None
        
    try:
        creds = service_account.Credentials.from_service_account_file(creds_path, scopes=SCOPES)
        return build('drive', 'v3', credentials=creds, cache_discovery=False)
    except Exception as e:
        logger.error(f"Failed to initialize Google Drive service from file: {e}")
        return None

def get_or_create_reports_folder(service, client_name: str) -> str:
    """Finds or creates a client-specific reports folder in Google Drive."""
    try:
        parent_id = os.getenv("GOOGLE_DRIVE_FOLDER_ID")
        if parent_id:
            query = f"name = 'Morning Tracker - {client_name}' and mimeType = 'application/vnd.google-apps.folder' and '{parent_id}' in parents and trashed = false"
        else:
            query = f"name = 'Morning Tracker - {client_name}' and mimeType = 'application/vnd.google-apps.folder' and trashed = false"
            
        results = service.files().list(
            q=query, 
            spaces='drive', 
            fields='files(id, name)',
            supportsAllDrives=True,
            includeItemsFromAllDrives=True
        ).execute()
        files = results.get('files', [])
        
        if files:
            return files[0]['id']
            
        folder_metadata = {
            'name': f'Morning Tracker - {client_name}',
            'mimeType': 'application/vnd.google-apps.folder'
        }
        if parent_id:
            folder_metadata['parents'] = [parent_id]
            
        folder = service.files().create(
            body=folder_metadata, 
            fields='id',
            supportsAllDrives=True
        ).execute()
        logger.info(f"Created new Google Drive folder: Morning Tracker - {client_name} (ID: {folder['id']})")
        return folder['id']
    except Exception as e:
        logger.error(f"Error getting/creating Google Drive folder: {e}")
        return None

def upload_docx_to_google_doc(docx_path: str, client_name: str, date_str: str, recipients: list, doc_suffix: str = "", template_path: str = None) -> str:
    """
    Uploads a local DOCX file to Google Drive and converts it to a native Google Doc.
    Supports continuous daily appending into a single monthly document (e.g. Scapia - June 2026).
    Appends new reports at the top of the monthly Google Doc to align with outline-navigation.
    Shares the Google Doc with the list of recipient email addresses.
    Returns the URL link of the Google Doc.
    """
    service = get_drive_service()
    if not service:
        return None
        
    try:
        folder_id = get_or_create_reports_folder(service, client_name)
        
        # Build Monthly Document Name (e.g., "Scapia - June 2026")
        month_year_str = datetime.now().strftime("%B %Y")
        file_name = f"{client_name}{doc_suffix} - {month_year_str}"
        
        # Search for an existing monthly Google Doc or docx inside the client's folder
        query = f"name = '{file_name}' and (mimeType = 'application/vnd.google-apps.document' or mimeType = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document') and '{folder_id}' in parents and trashed = false"
        results = service.files().list(
            q=query, 
            spaces='drive', 
            fields='files(id, webViewLink, mimeType)',
            supportsAllDrives=True,
            includeItemsFromAllDrives=True
        ).execute()
        files = results.get('files', [])
        
        doc_id = None
        web_link = None
        
        if files:
            # Monthly document already exists -> Export, Merge, and Update in-place
            doc_id = files[0]['id']
            web_link = files[0]['webViewLink']
            mime_type = files[0].get('mimeType')
            logger.info(f"Monthly document exists: ID {doc_id} ({mime_type}). Initiating daily append merge...")
            
            temp_existing_path = docx_path + ".existing.docx"
            temp_combined_path = docx_path + ".combined.docx"
            
            media = None
            try:
                # 1. Export Google Doc or download standard binary file using downloader
                from googleapiclient.http import MediaIoBaseDownload
                import io
                
                if mime_type == 'application/vnd.google-apps.document':
                    export_request = service.files().export_media(
                        fileId=doc_id, 
                        mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )
                else:
                    export_request = service.files().get_media(
                        fileId=doc_id
                    )
                fh = io.BytesIO()
                downloader = MediaIoBaseDownload(fh, export_request)
                done = False
                while not done:
                    status, done = downloader.next_chunk()
                    
                with open(temp_existing_path, "wb") as f:
                    f.write(fh.getvalue())
                
                # 2. Merge new daily report with existing month report (new on top)
                merge_docx_files(docx_path, temp_existing_path, temp_combined_path, template_path=template_path)
                
                # 3. Update the existing Google Doc content
                media = MediaFileUpload(
                    temp_combined_path,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    resumable=True
                )

                service.files().update(
                    fileId=doc_id,
                    media_body=media,
                    supportsAllDrives=True
                ).execute()
                
                logger.info(f"Successfully appended report to monthly Google Doc: {file_name}")
                
            except Exception as merge_err:
                logger.error(f"Failed to merge with existing monthly doc: {merge_err}", exc_info=True)
            finally:
                # Safely release Windows file lock
                if media and hasattr(media, '_fd') and media._fd:
                    try: media._fd.close()
                    except: pass
                # Cleanup temp files
                if os.path.exists(temp_existing_path):
                    os.remove(temp_existing_path)
                if os.path.exists(temp_combined_path):
                    os.remove(temp_combined_path)
        
        if not doc_id:
            # Monthly document does not exist yet -> Create new
            logger.info(f"Creating new monthly Google Doc: {file_name}")
            file_metadata = {
                'name': file_name,
                'mimeType': 'application/vnd.google-apps.document'
            }
            if folder_id:
                file_metadata['parents'] = [folder_id]
                
            media = None
            try:
                media = MediaFileUpload(
                    docx_path, 
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                    resumable=True
                )
                uploaded_file = service.files().create(
                    body=file_metadata, 
                    media_body=media, 
                    fields='id, webViewLink',
                    supportsAllDrives=True
                ).execute()
                
                doc_id = uploaded_file.get('id')
                web_link = uploaded_file.get('webViewLink')
                logger.info(f"Created monthly Google Doc ID: {doc_id}")
            finally:
                if media and hasattr(media, '_fd') and media._fd:
                    try: media._fd.close()
                    except: pass
            
        # Share document with recipient emails (always run to ensure new recipients get access)
        for email in recipients:
            if not email:
                continue
            try:
                user_permission = {
                    'type': 'user',
                    'role': 'writer',
                    'emailAddress': email
                }
                service.permissions().create(
                    fileId=doc_id, 
                    body=user_permission, 
                    sendNotificationEmails=False,
                    supportsAllDrives=True
                ).execute()
            except Exception as share_err:
                # Ignore duplicate sharing errors or domain restriction warnings
                pass
                
        # Make document publicly readable via link (Anyone can view)
        try:
            public_permission = {
                'type': 'anyone',
                'role': 'reader'
            }
            service.permissions().create(
                fileId=doc_id,
                body=public_permission,
                supportsAllDrives=True
            ).execute()
            logger.info(f"Configured public sharing permission for Google Doc: {doc_id}")
        except Exception as public_err:
            logger.error(f"Failed to set public sharing permission for Google Doc {doc_id}: {public_err}")
                
        return web_link
        
    except Exception as e:
        logger.error(f"Failed to upload report to Google Docs: {e}", exc_info=True)
        return None

def write_articles_to_sheet(ws, date_str, client_name, grouped_data):
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    
    # Style definitions (Steel Blue & Ice Blue theme)
    master_fill = PatternFill(start_color="365F91", end_color="365F91", fill_type="solid")
    master_font = Font(name="Calibri", size=12, bold=True, color="FFFFFF")
    
    table_header_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
    table_header_font = Font(name="Calibri", size=10, bold=True, color="000000")
    
    data_font = Font(name="Calibri", size=10)
    link_font = Font(name="Calibri", size=10, color="0000FF", underline="single")
    title_font = Font(name="Calibri", size=14, bold=True, color="365F91")
    subtitle_font = Font(name="Calibri", size=11, italic=True)

    thin_border = Border(
        left=Side(style='thin', color='D3D3D3'),
        right=Side(style='thin', color='D3D3D3'),
        top=Side(style='thin', color='D3D3D3'),
        bottom=Side(style='thin', color='D3D3D3')
    )

    max_cols = 6 # Link, Title, Author, Publication, Date, Summary

    # Title Block
    ws.cell(row=1, column=1, value="NEXUS NEWS BRIEFING: CUMULATIVE REPORT").font = title_font
    ws.cell(row=2, column=1, value=f"Date: {date_str} | Generated for {client_name}").font = subtitle_font
    ws.row_dimensions[1].height = 25
    ws.row_dimensions[2].height = 20

    current_row = 4

    headers = [
        "Link of the article",
        "Title of the article",
        "Author of the article",
        "Name of publication",
        "Time of publishing",
        "Summary of the article"
    ]

    # Iterate through sections
    for section_name, articles in grouped_data.items():
        if not articles:
            continue

        # Add Section Heading Row (Merged A to F)
        ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=max_cols)
        c_master = ws.cell(row=current_row, column=1, value=section_name)
        c_master.font = master_font
        c_master.fill = master_fill
        c_master.alignment = Alignment(vertical="center", indent=1)
        ws.row_dimensions[current_row].height = 26
        current_row += 1

        # Add Table Headers for the nested table
        for col_idx, h in enumerate(headers, start=1):
            cell = ws.cell(row=current_row, column=col_idx, value=h)
            cell.fill = table_header_fill
            cell.font = table_header_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = thin_border
        ws.row_dimensions[current_row].height = 20
        current_row += 1

        # Add articles
        for art in articles:
            title = art.get("title", "Untitled Article")
            url = art.get("url") or art.get("resolved_url") or "#"
            author = art.get("author") or "N/A"
            pub_name = art.get("agency") or "Unknown Publication"
            pub_date = art.get("published_at")
            
            date_text = "N/A"
            if pub_date:
                if isinstance(pub_date, str):
                    date_text = pub_date[:10]
                else:
                    date_text = pub_date.strftime("%Y-%m-%d")

            # Write values
            c_link = ws.cell(row=current_row, column=1, value="Link")
            c_link.hyperlink = url
            c_link.font = link_font
            c_link.alignment = Alignment(horizontal="center")
            
            c_title = ws.cell(row=current_row, column=2, value=title)
            
            author_val = str(author).strip()
            if not author_val or author_val.upper() in ("N/A", "NONE", "NULL", ""):
                author_val = "N/A"
            c_author = ws.cell(row=current_row, column=3, value=author_val)
            
            c_pub = ws.cell(row=current_row, column=4, value=pub_name)
            c_date = ws.cell(row=current_row, column=5, value=date_text)
            
            summary_text = art.get("summary") or art.get("_summary") or art.get("full_body") or ""
            c_summary = ws.cell(row=current_row, column=6, value=summary_text)

            # Borders and alignment
            for col_idx in range(1, max_cols + 1):
                c = ws.cell(row=current_row, column=col_idx)
                c.border = thin_border
                c.font = data_font
                if col_idx == 1:
                    pass
                elif col_idx == max_cols:
                    c.alignment = Alignment(vertical="center", wrap_text=True)
                else:
                    c.alignment = Alignment(vertical="center", wrap_text=False)

            ws.row_dimensions[current_row].height = 20
            current_row += 1

        # Empty row for breathing space
        current_row += 1

    # Set column widths
    ws.column_dimensions["A"].width = 15  # Link
    ws.column_dimensions["B"].width = 50  # Title
    ws.column_dimensions["C"].width = 25  # Author
    ws.column_dimensions["D"].width = 25  # Publication Name
    ws.column_dimensions["E"].width = 15  # Date
    ws.column_dimensions["F"].width = 60  # Summary

def upload_cumulative_excel_to_google_sheet(grouped_data: dict, client_name: str, date_str: str, recipients: list) -> str:
    """
    Creates or updates a cumulative Excel sheet stored on Google Drive (.xlsx format, openable via Google Sheets).
    Daily run articles are added into a separate tab named after the execution date.
    Shares the sheet with the list of recipient emails.
    Returns the URL link to the sheet.
    """
    service = get_drive_service()
    if not service:
        return None
        
    try:
        folder_id = get_or_create_reports_folder(service, client_name)
        file_name = f"{client_name} - Cumulative Filtered Articles"
        
        # Search for an existing cumulative sheet inside the client's folder
        query = f"name = '{file_name}' and mimeType = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet' and '{folder_id}' in parents and trashed = false"
        results = service.files().list(
            q=query, 
            spaces='drive', 
            fields='files(id, webViewLink)',
            supportsAllDrives=True,
            includeItemsFromAllDrives=True
        ).execute()
        files = results.get('files', [])
        
        sheet_id = None
        web_link = None
        
        import openpyxl
        temp_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "reports")
        os.makedirs(temp_dir, exist_ok=True)
        local_path = os.path.join(temp_dir, f"Cumulative_{client_name.replace(' ', '_')}.xlsx")
        
        if files:
            sheet_id = files[0]['id']
            web_link = files[0]['webViewLink']
            logger.info(f"Cumulative sheet exists: ID {sheet_id}. Downloading for daily tab append...")
            
            temp_existing_path = local_path + ".existing.xlsx"
            
            media = None
            try:
                # 1. Download standard binary excel file
                from googleapiclient.http import MediaIoBaseDownload
                import io
                
                request = service.files().get_media(fileId=sheet_id)
                fh = io.BytesIO()
                downloader = MediaIoBaseDownload(fh, request)
                done = False
                while not done:
                    status, done = downloader.next_chunk()
                    
                with open(temp_existing_path, "wb") as f:
                    f.write(fh.getvalue())
                
                # 2. Open and add/overwrite today's date tab
                wb = openpyxl.load_workbook(temp_existing_path)
                
                # Remove if tab already exists to avoid duplicate tabs for the same date
                if date_str in wb.sheetnames:
                    wb.remove(wb[date_str])
                    
                # Create sheet tab (position 0 to show newest date first)
                ws = wb.create_sheet(title=date_str, index=0)
                write_articles_to_sheet(ws, date_str, client_name, grouped_data)
                
                wb.save(local_path)
                
                # 3. Update the existing file content on Google Drive
                media = MediaFileUpload(
                    local_path, 
                    mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 
                    resumable=True
                )
                service.files().update(
                    fileId=sheet_id,
                    media_body=media,
                    supportsAllDrives=True
                ).execute()
                logger.info(f"Successfully appended today's tab to cumulative sheet: {file_name}")
                
            except Exception as merge_err:
                logger.error(f"Failed to merge with existing cumulative sheet: {merge_err}", exc_info=True)
            finally:
                if media and hasattr(media, '_fd') and media._fd:
                    try: media._fd.close()
                    except: pass
                if os.path.exists(temp_existing_path):
                    try: os.remove(temp_existing_path)
                    except: pass
                if os.path.exists(local_path):
                    try: os.remove(local_path)
                    except: pass
        
        if not sheet_id:
            logger.info(f"Creating new cumulative spreadsheet: {file_name}")
            wb = openpyxl.Workbook()
            # Rename default sheet
            ws = wb.active
            ws.title = date_str
            write_articles_to_sheet(ws, date_str, client_name, grouped_data)
            
            wb.save(local_path)
            
            file_metadata = {
                'name': file_name,
                'mimeType': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            }
            if folder_id:
                file_metadata['parents'] = [folder_id]
                
            media = None
            try:
                media = MediaFileUpload(
                    local_path, 
                    mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 
                    resumable=True
                )
                uploaded_file = service.files().create(
                    body=file_metadata, 
                    media_body=media, 
                    fields='id, webViewLink',
                    supportsAllDrives=True
                ).execute()
                
                sheet_id = uploaded_file.get('id')
                web_link = uploaded_file.get('webViewLink')
                logger.info(f"Created cumulative sheet ID: {sheet_id}")
            finally:
                if media and hasattr(media, '_fd') and media._fd:
                    try: media._fd.close()
                    except: pass
                if os.path.exists(local_path):
                    try: os.remove(local_path)
                    except: pass
                    
        # Share spreadsheet with recipients
        for email in recipients:
            if not email:
                continue
            try:
                user_permission = {
                    'type': 'user',
                    'role': 'writer',
                    'emailAddress': email
                }
                service.permissions().create(
                    fileId=sheet_id, 
                    body=user_permission, 
                    sendNotificationEmails=False,
                    supportsAllDrives=True
                ).execute()
            except Exception:
                pass
                
        # Make document publicly readable via link (Anyone can view)
        try:
            public_permission = {
                'type': 'anyone',
                'role': 'reader'
            }
            service.permissions().create(
                fileId=sheet_id,
                body=public_permission,
                supportsAllDrives=True
            ).execute()
        except Exception:
            pass
            
        return web_link
        
    except Exception as e:
        logger.error(f"Failed to upload cumulative sheet to Google Drive: {e}", exc_info=True)
        return None


def append_daily_takeaways_to_sheet(company_name: str, run_date: Any, takeaways_text: str) -> Optional[str]:
    """
    Appends daily strategic takeaways to a dedicated Excel spreadsheet on Google Drive,
    creating/syncing it in place, using separate tabs for each month (e.g. 'July 2026').
    Does NOT email it automatically. Returns the web link to the Google Sheet.
    """
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    
    if not takeaways_text:
        logger.warning(f"[Takeaways Sheet] No takeaways to append for {company_name}")
        return None
        
    service = get_drive_service()
    if not service:
        logger.error("[Takeaways Sheet] Google Drive service not initialized.")
        return None
        
    local_path = ""
    try:
        folder_id = get_or_create_reports_folder(service, company_name)
        file_name = f"{company_name} - Strategic Takeaways History"
        
        # 1. Search for existing spreadsheet
        query = f"name = '{file_name}' and mimeType = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet' and '{folder_id}' in parents and trashed = false"
        results = service.files().list(
            q=query, 
            spaces='drive', 
            fields='files(id, webViewLink)',
            supportsAllDrives=True,
            includeItemsFromAllDrives=True
        ).execute()
        files = results.get('files', [])
        
        sheet_id = None
        web_link = None
        
        temp_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "reports")
        os.makedirs(temp_dir, exist_ok=True)
        local_path = os.path.join(temp_dir, f"Takeaways_{company_name.replace(' ', '_')}.xlsx")
        
        if files:
            sheet_id = files[0]['id']
            web_link = files[0]['webViewLink']
            logger.info(f"[Takeaways Sheet] Sheet exists (ID: {sheet_id}). Downloading...")
            # Download file
            from googleapiclient.http import MediaIoBaseDownload
            import io
            request = service.files().get_media(fileId=sheet_id)
            fh = io.FileIO(local_path, 'wb')
            downloader = MediaIoBaseDownload(fh, request)
            done = False
            while not done:
                _, done = downloader.next_chunk()
            fh.close()
            wb = openpyxl.load_workbook(local_path)
        else:
            logger.info(f"[Takeaways Sheet] Creating new sheet...")
            wb = openpyxl.Workbook()
            # remove default sheet
            default_sheet = wb.active
            if default_sheet:
                wb.remove(default_sheet)
                
        # 2. Determine month tab name
        if isinstance(run_date, str):
            try:
                date_obj = datetime.strptime(run_date, "%Y-%m-%d").date()
            except Exception:
                try:
                    date_obj = datetime.strptime(run_date, "%d %B %Y").date()
                except Exception:
                    date_obj = datetime.today().date()
        elif isinstance(run_date, (datetime, date)):
            date_obj = run_date
        else:
            date_obj = datetime.today().date()
            
        month_tab_name = date_obj.strftime("%B %Y")  # e.g., "July 2026"
        date_str_iso = date_obj.strftime("%Y-%m-%d")
        
        # 3. Get or create month sheet
        if month_tab_name in wb.sheetnames:
            ws = wb[month_tab_name]
        else:
            ws = wb.create_sheet(title=month_tab_name)
            # Write header
            header_font = Font(name="Calibri", size=11, bold=True, color="ffffff")
            header_fill = PatternFill(start_color="4285F4", end_color="4285F4", fill_type="solid")  # Google Blue
            header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
            
            headers = ["Date", "Takeaway Title", "Details"]
            for col_idx, header_text in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col_idx, value=header_text)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_align
                
            ws.row_dimensions[1].height = 25
            
        # 4. Parse daily takeaways text and append
        lines = [l.strip().lstrip("-*•").strip() for l in takeaways_text.split("\n") if l.strip()]
        
        data_font = Font(name="Calibri", size=11)
        align_left = Alignment(horizontal="left", vertical="top", wrap_text=True)
        align_center = Alignment(horizontal="center", vertical="top")
        
        thin_side = Side(border_style="thin", color="D3D3D3")
        thin_border = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)
        
        start_row = ws.max_row + 1
        
        added_any = False
        for line in lines:
            # Match titles split by "—" or ":"
            parts = line.split("—", 1)
            if len(parts) == 2:
                title_part, text_part = parts[0].strip(), parts[1].strip()
            else:
                parts = line.split(":", 1)
                if len(parts) == 2:
                    title_part, text_part = parts[0].strip(), parts[1].strip()
                else:
                    title_part = "Key Insight"
                    text_part = line.strip()
                    
            if not text_part:
                continue
                
            # Append row
            c_date = ws.cell(row=start_row, column=1, value=date_str_iso)
            c_title = ws.cell(row=start_row, column=2, value=title_part)
            c_text = ws.cell(row=start_row, column=3, value=text_part)
            
            # Formatting
            for cell in (c_date, c_title, c_text):
                cell.font = data_font
                cell.border = thin_border
                
            c_date.alignment = align_center
            c_title.alignment = align_left
            c_text.alignment = align_left
            
            start_row += 1
            added_any = True
            
        if not added_any:
            # Fallback
            c_date = ws.cell(row=start_row, column=1, value=date_str_iso)
            c_title = ws.cell(row=start_row, column=2, value="Daily Takeaways Summary")
            c_text = ws.cell(row=start_row, column=3, value=takeaways_text)
            
            for cell in (c_date, c_title, c_text):
                cell.font = data_font
                cell.border = thin_border
            c_date.alignment = align_center
            c_title.alignment = align_left
            c_text.alignment = align_left
            
        # Adjust column widths
        ws.column_dimensions["A"].width = 15
        ws.column_dimensions["B"].width = 30
        ws.column_dimensions["C"].width = 80
        
        # Save workbook locally
        wb.save(local_path)
        
        # 5. Upload back to Google Drive
        from googleapiclient.http import MediaFileUpload
        media = MediaFileUpload(local_path, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', resumable=True)
        
        if sheet_id:
            # Update existing file
            service.files().update(
                fileId=sheet_id,
                media_body=media,
                supportsAllDrives=True
            ).execute()
        else:
            # Create new file
            file_metadata = {
                'name': file_name,
                'parents': [folder_id],
                'mimeType': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            }
            file = service.files().create(
                body=file_metadata,
                media_body=media,
                fields='id, webViewLink',
                supportsAllDrives=True
            ).execute()
            sheet_id = file.get('id')
            web_link = file.get('webViewLink')
            
        # Ensure public link sharing is enabled (anyone with link can read)
        try:
            public_permission = {
                'type': 'anyone',
                'role': 'reader'
            }
            service.permissions().create(
                fileId=sheet_id,
                body=public_permission,
                supportsAllDrives=True
            ).execute()
        except Exception:
            pass
            
        # Clean up local file safely
        try:
            if local_path and os.path.exists(local_path):
                os.remove(local_path)
        except Exception:
            pass
            
        return web_link
        
    except Exception as e:
        logger.error(f"[Takeaways Sheet] Error updating Google Sheet for {company_name}: {e}", exc_info=True)
        if local_path and os.path.exists(local_path):
            try:
                os.remove(local_path)
            except Exception:
                pass
        return None


def download_takeaways_sheet_file(company_name: str) -> Optional[str]:
    """
    Searches for the takeaways history Excel sheet for the specified company on Google Drive,
    downloads it locally to a temp path, and returns the path.
    """
    service = get_drive_service()
    if not service:
        logger.error("[Takeaways Sheet] Google Drive service not initialized.")
        return None
        
    try:
        folder_id = get_or_create_reports_folder(service, company_name)
        file_name = f"{company_name} - Strategic Takeaways History"
        
        query = f"name = '{file_name}' and mimeType = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet' and '{folder_id}' in parents and trashed = false"
        results = service.files().list(
            q=query, 
            spaces='drive', 
            fields='files(id)',
            supportsAllDrives=True,
            includeItemsFromAllDrives=True
        ).execute()
        files = results.get('files', [])
        
        if not files:
            logger.warning(f"[Takeaways Sheet] Takeaways sheet not found for {company_name}")
            return None
            
        sheet_id = files[0]['id']
        
        temp_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "reports")
        os.makedirs(temp_dir, exist_ok=True)
        local_path = os.path.join(temp_dir, f"Takeaways_{company_name.replace(' ', '_')}_Download.xlsx")
        
        from googleapiclient.http import MediaIoBaseDownload
        import io
        request = service.files().get_media(fileId=sheet_id)
        fh = io.FileIO(local_path, 'wb')
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            _, done = downloader.next_chunk()
        fh.close()
        
        return local_path
    except Exception as e:
        logger.error(f"[Takeaways Sheet] Failed to download takeaways sheet for {company_name}: {e}", exc_info=True)
        return None
